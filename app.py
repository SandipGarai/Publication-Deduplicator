"""
Publication Deduplicator v1.0
"""

import streamlit as st
import re
import io
from collections import Counter, defaultdict
from difflib import SequenceMatcher
from typing import List, Tuple, Dict, Optional
import hashlib
import urllib.parse
import requests
import time


# Page configuration
st.set_page_config(
    page_title="Publication Deduplicator",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1200px;
    }

    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        color: #ffffff !important;
        text-align: center;
    }

    .main-header h1 {
        margin: 0;
        font-size: 2.5rem;
        font-weight: 700;
        color: #ffffff !important;
    }

    .main-header p {
        margin: 0.5rem 0 0 0;
        opacity: 0.95;
        font-size: 1.1rem;
        color: #ffffff !important;
    }

    .metric-card {
        background: #ffffff !important;
        border-radius: 12px;
        padding: 1.5rem;
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
        border-left: 5px solid;
        margin-bottom: 1rem;
    }

    .metric-card.total { border-left-color: #667eea; }
    .metric-card.unique { border-left-color: #28a745; }
    .metric-card.duplicate { border-left-color: #dc3545; }
    .metric-card.reduction { border-left-color: #ffc107; }

    .metric-value {
        font-size: 2.5rem;
        font-weight: 700;
        margin: 0;
        line-height: 1.2;
        color: #1a1a2e !important;
    }

    .metric-label {
        font-size: 0.85rem;
        color: #666666 !important;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        margin-top: 0.25rem;
        font-weight: 500;
    }

    .article-card {
        background: #ffffff !important;
        border-radius: 8px;
        padding: 1rem 1.25rem;
        margin-bottom: 0.75rem;
        border-left: 4px solid #667eea;
        font-size: 0.95rem;
        line-height: 1.6;
        color: #222222 !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }

    .article-card.duplicate {
        border-left-color: #dc3545;
        background: #fff8f8 !important;
    }

    .article-card strong {
        color: #1a1a2e !important;
    }

    .rating-badge {
        display: inline-block;
        background: #17a2b8;
        color: #ffffff !important;
        padding: 0.15rem 0.5rem;
        border-radius: 10px;
        font-size: 0.7rem;
        font-weight: 600;
        margin-left: 0.3rem;
    }

    .rating-badge.naas { background: #28a745; }
    .rating-badge.if { background: #fd7e14; }

    .crossref-link {
        display: inline-block;
        background: #0066cc;
        color: #ffffff !important;
        padding: 0.2rem 0.5rem;
        border-radius: 6px;
        font-size: 0.75rem;
        font-weight: 500;
        text-decoration: none;
        margin-top: 0.5rem;
    }

    .crossref-link:hover {
        background: #0052a3;
        color: #ffffff !important;
    }

    .section-header {
        display: flex;
        align-items: center;
        gap: 0.75rem;
        margin: 2rem 0 1rem 0;
        padding-bottom: 0.75rem;
        border-bottom: 2px solid #555555;
    }

    .section-header h2 {
        margin: 0;
        font-size: 1.5rem;
        font-weight: 600;
    }

    .info-box {
        background: #e3f2fd !important;
        border: 1px solid #90caf9;
        border-radius: 8px;
        padding: 1rem 1.25rem;
        margin: 1rem 0;
        color: #1a237e !important;
    }

    .info-box.warning {
        background: #fff8e1 !important;
        border-color: #ffca28;
        color: #5d4037 !important;
    }

    .info-box.success {
        background: #e8f5e9 !important;
        border-color: #81c784;
        color: #1b5e20 !important;
    }

    .info-box.error {
        background: #ffebee !important;
        border-color: #ef9a9a;
        color: #b71c1c !important;
    }

    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)


def extract_doi(article: str) -> Optional[str]:
    """Extract DOI from article text."""
    # Multiple DOI patterns
    patterns = [
        r'https?://doi\.org/(\d{2}\.\d+/[^\s,\)\]\>]+)',
        r'doi[:\s]+(\d{2}\.\d+/[^\s,\)\]\>]+)',
        r'DOI[:\s]+(\d{2}\.\d+/[^\s,\)\]\>]+)',
        r'(\d{2}\.\d{4,}/[^\s,\)\]\>]+)',
    ]

    for pattern in patterns:
        match = re.search(pattern, article, re.IGNORECASE)
        if match:
            doi = match.group(1).rstrip('.').rstrip(
                ')').rstrip(']').rstrip('>')
            return doi
    return None


def remove_doi_from_text(article: str) -> str:
    """
    Remove DOI or DOI URL from article text for clean display.
    """
    patterns = [
        r'https?://doi\.org/\S+',
        r'\bdoi:\s*\S+',
        r'\bDOI:\s*\S+',
        r'\b\d{2}\.\d{4,}/\S+'
    ]

    cleaned = article
    for pattern in patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)

    cleaned = re.sub(r'\s{2,}', ' ', cleaned).strip()
    return cleaned


@st.cache_data(show_spinner=False)
def fetch_doi_from_crossref(article: str) -> Optional[str]:
    """
    Fetch DOI from Crossref for articles without DOI.
    Uses title + year and validates via similarity.
    """

    # ---- Extract title ----
    title_match = re.search(r'\(\d{4}\)[.,]?\s*([^.]+)\.', article)
    if not title_match:
        return None

    title = title_match.group(1).strip()

    # ---- Extract year ----
    year_match = re.search(r'\((\d{4})\)', article)
    year = year_match.group(1) if year_match else None

    params = {
        "query.title": title,
        "rows": 3,
        "mailto": "your_email@domain.com"  # recommended by Crossref
    }

    if year:
        params["filter"] = f"from-pub-date:{year},until-pub-date:{year}"

    try:
        response = requests.get(
            "https://api.crossref.org/works",
            params=params,
            timeout=10
        )

        if response.status_code != 200:
            return None

        items = response.json()["message"]["items"]
        if not items:
            return None

        best = items[0]
        candidate_title = best.get("title", [""])[0]

        similarity = calculate_similarity(
            normalize_text(title),
            normalize_text(candidate_title)
        )

        if similarity >= 0.85:
            return best.get("DOI")

    except Exception:
        return None

    finally:
        time.sleep(0.5)  # polite rate limiting

    return None


def generate_crossref_link(article: str, auto_fetch: bool = True) -> Optional[str]:
    """
    Generate CrossRef/DOI link.
    Uses existing DOI first, otherwise fetches from Crossref.
    """
    doi = extract_doi(article)

    if not doi and auto_fetch:
        doi = fetch_doi_from_crossref(article)

    if doi:
        return f"https://doi.org/{doi}"

    return None


def extract_rating(article: str) -> Dict[str, str]:
    """Extract NAAS rating and Impact Factor."""
    ratings = {}

    naas_patterns = [
        r'\(NAAS\s+rating\s+([\d.]+)\)',
        r'\(NAAS\s+([\d.]+)\)',
        r'NAAS\s+rating\s+([\d.]+)',
        r'NAAS\s+([\d.]+)',
    ]

    for pattern in naas_patterns:
        match = re.search(pattern, article, re.IGNORECASE)
        if match:
            ratings['naas'] = match.group(1)
            break

    if_patterns = [
        r'SCI[/-]?Clarivate[:\s]*([\d.]+)',
        r'Thomson\s+Reuters\s+IF[:\s]*([\d.]+)',
        r'\(IF[:\s]*([\d.]+)\)',
    ]

    for pattern in if_patterns:
        match = re.search(pattern, article, re.IGNORECASE)
        if match:
            ratings['if'] = match.group(1)
            break

    return ratings


def remove_naas_from_text(article: str) -> str:
    """
    Remove any NAAS-related text from article for uniform display.
    """
    patterns = [
        r'\(NAAS\s+rating\s+[\d.]+\)',
        r'\(NAAS\s+[\d.]+\)',
        r'NAAS\s+rating\s+[\d.]+',
        r'NAAS\s+[\d.]+',
    ]

    cleaned = article
    for pattern in patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)

    # Clean extra spaces
    cleaned = re.sub(r'\s{2,}', ' ', cleaned).strip()
    return cleaned


def normalize_text(text: str) -> str:
    """Normalize text for comparison."""
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[–—−]', '-', text)
    text = re.sub(r'https?://[^\s]+', '', text)
    text = re.sub(r'doi:\s*\S+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\(naas[^)]*\)', '', text, flags=re.IGNORECASE)
    text = re.sub(r'naas\s+[\d.]+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'published\s+on[^.]+', '', text, flags=re.IGNORECASE)
    text = text.strip()
    return text


def generate_article_hash(article: str) -> str:
    """Generate hash for duplicate detection."""
    return hashlib.md5(normalize_text(article).encode()).hexdigest()


def calculate_similarity(text1: str, text2: str) -> float:
    """Calculate similarity between texts."""
    return SequenceMatcher(None, text1, text2).ratio()


def detect_format(text: str) -> str:
    """Detect the format of the input text."""
    # Count numbered lines
    numbered_lines = len(re.findall(r'^\d+\.\s', text, re.MULTILINE))

    # Count paragraph breaks
    paragraph_breaks = len(re.findall(r'\n\s*\n', text))

    # Count author patterns
    author_patterns = len(re.findall(r'[A-Z][a-z]+,\s+[A-Z]\.', text))

    if numbered_lines >= 5:
        return "numbered"
    elif paragraph_breaks >= 5:
        return "paragraphs"
    else:
        return "mixed"


def parse_numbered_articles(text: str) -> List[Tuple[int, str]]:
    """Parse strictly numbered articles."""
    articles = []

    # Split by number at start of line
    pattern = r'(?:^|\n)(\d+)\.\s+'
    parts = re.split(pattern, text)

    # parts: [before_first, '1', content1, '2', content2, ...]
    i = 1
    while i < len(parts) - 1:
        try:
            num = int(parts[i])
            content = parts[i + 1].strip()
            if content:
                # Clean multi-line content
                content = ' '.join(content.split())
                articles.append((num, content))
        except (ValueError, IndexError):
            pass
        i += 2

    return articles


def parse_paragraph_articles(text: str) -> List[str]:
    """Parse articles separated by blank lines."""
    articles = []
    blocks = re.split(r'\n\s*\n+', text)

    for block in blocks:
        block = block.strip()
        if not block or len(block) < 40:
            continue

        # Clean up
        block = re.sub(r'^[\s>]+', '', block, flags=re.MULTILINE)
        block = ' '.join(block.split())
        block = re.sub(r'^\d+[\.\)]\s*', '', block)

        # Validate it looks like an article
        has_year = bool(re.search(r'\(\d{4}\)', block)) or bool(
            re.search(r'\b\d{4}\b', block))
        has_author = bool(re.search(r'[A-Z][a-z]+,?\s+[A-Z]\.', block))

        if has_year or has_author:
            articles.append(block)

    return articles


def split_concatenated(text: str) -> str:
    """Split articles concatenated without line breaks."""
    # DOI/URL followed by author
    text = re.sub(
        r'(https?://[^\s]+?)([A-Z][a-z]+,\s+[A-Z]\.)', r'\1\n\n\2', text)
    # Period + author pattern
    text = re.sub(
        r'\.(\s*)([A-Z][a-z]{2,},\s+[A-Z]\.\s*[A-Z]?)', r'.\n\n\2', text)
    # Year followed by author
    text = re.sub(
        r'(\(\d{4}\)[.,]?)([A-Z][a-z]+,\s+[A-Z]\.)', r'\1\n\n\2', text)
    return text


def parse_articles(text: str, trust_numbering: bool = True) -> Tuple[List[str], str, Dict]:
    """
    Parse articles with format detection.

    Returns:
        - List of articles
        - Detected format
        - Parsing info dict
    """
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Detect format
    format_type = detect_format(text)
    info = {'format': format_type, 'numbered_count': 0, 'paragraph_count': 0}

    # Try numbered parsing first if format suggests it
    if format_type == "numbered" and trust_numbering:
        numbered_articles = parse_numbered_articles(text)
        info['numbered_count'] = len(numbered_articles)

        if numbered_articles:
            # Check for gaps in numbering
            numbers = [n for n, _ in numbered_articles]
            expected = list(range(1, max(numbers) + 1)) if numbers else []
            missing = set(expected) - set(numbers)
            info['missing_numbers'] = sorted(missing) if missing else []
            info['max_number'] = max(numbers) if numbers else 0

            return [art for _, art in numbered_articles], format_type, info

    # Fall back to paragraph parsing
    text = split_concatenated(text)
    articles = parse_paragraph_articles(text)
    info['paragraph_count'] = len(articles)

    return articles, format_type, info


def find_duplicates(articles: List[str], similarity_threshold: float = 0.85) -> Tuple[List[str], Dict[int, List[int]]]:
    """Find and remove duplicate articles."""
    article_occurrences = defaultdict(list)
    doi_to_indices = defaultdict(list)

    # First pass: hash and DOI matching
    for idx, article in enumerate(articles):
        article_hash = generate_article_hash(article)
        doi = extract_doi(article)

        if doi:
            doi_to_indices[doi.lower()].append(idx)

        article_occurrences[article_hash].append(idx)

    # Merge DOI-based duplicates
    for doi, indices in doi_to_indices.items():
        if len(indices) > 1:
            first_hash = generate_article_hash(articles[indices[0]])
            for idx in indices[1:]:
                idx_hash = generate_article_hash(articles[idx])
                if idx_hash != first_hash:
                    if idx in article_occurrences[idx_hash]:
                        article_occurrences[idx_hash].remove(idx)
                    if idx not in article_occurrences[first_hash]:
                        article_occurrences[first_hash].append(idx)

    # Second pass: fuzzy matching
    unique_hashes = [
        h for h, indices in article_occurrences.items() if indices]

    for i, hash1 in enumerate(unique_hashes):
        if not article_occurrences[hash1]:
            continue
        idx1 = article_occurrences[hash1][0]
        norm1 = normalize_text(articles[idx1])

        for hash2 in unique_hashes[i+1:]:
            if not article_occurrences[hash2]:
                continue
            idx2 = article_occurrences[hash2][0]
            norm2 = normalize_text(articles[idx2])

            if calculate_similarity(norm1, norm2) >= similarity_threshold:
                article_occurrences[hash1].extend(article_occurrences[hash2])
                article_occurrences[hash2] = []

    # Build results
    unique_articles = []
    duplicate_groups = {}
    processed = set()

    for article_hash, indices in article_occurrences.items():
        if not indices:
            continue

        indices = sorted(set(indices))
        first_idx = indices[0]

        if first_idx not in processed:
            unique_articles.append(articles[first_idx])
            processed.add(first_idx)

            if len(indices) > 1:
                duplicate_groups[first_idx] = indices[1:]
                for idx in indices[1:]:
                    processed.add(idx)

    return unique_articles, duplicate_groups


def format_article_html(
    article: str,
    show_crossref: bool = True,
    auto_fetch_doi: bool = True
) -> str:
    """Format article with badges and CrossRef link."""
    ratings = extract_rating(article)
    badges = ""

    if 'naas' in ratings:
        badges += f' <span class="rating-badge naas">NAAS: {ratings["naas"]}</span>'
    if 'if' in ratings:
        badges += f' <span class="rating-badge if">IF: {ratings["if"]}</span>'

    crossref = ""
    if show_crossref:
        link = generate_crossref_link(article, auto_fetch=auto_fetch_doi)

        if link:
            crossref = f'<br/><a href="{link}" target="_blank" class="crossref-link">🔗 View on CrossRef/DOI</a>'

    clean_article = remove_doi_from_text(remove_naas_from_text(article))

    return clean_article + badges + crossref


def create_docx(articles: List[str]) -> bytes:
    """Create DOCX with articles and CrossRef links."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading('Unique Articles', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Total unique articles: {len(articles)}')
        doc.add_paragraph('')

        for idx, article in enumerate(articles, 1):
            ratings = extract_rating(article)
            link = generate_crossref_link(article, auto_fetch=True)

            rating_str = ""
            if 'naas' in ratings:
                rating_str += f" [NAAS: {ratings['naas']}]"
            if 'if' in ratings:
                rating_str += f" [IF: {ratings['if']}]"

            para = doc.add_paragraph()
            clean_article = remove_doi_from_text(
                remove_naas_from_text(article))

            run = para.add_run(f"{clean_article}{rating_str}")

            run.font.size = Pt(11)

            if link:
                para.add_run(f"\n   CrossRef: {link}")

            para.paragraph_format.space_after = Pt(12)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except ImportError:
        return None


def create_text_output(articles: List[str]) -> str:
    """Create clean text output with CrossRef links (NO numbering)."""
    lines = [
        "UNIQUE RESEARCH ARTICLES",
        "=" * 60,
        f"Total: {len(articles)} articles",
        ""
    ]

    for article in articles:
        ratings = extract_rating(article)
        link = generate_crossref_link(article, auto_fetch=True)

        rating_str = ""
        if 'naas' in ratings:
            rating_str += f" [NAAS: {ratings['naas']}]"
        if 'if' in ratings:
            rating_str += f" [IF: {ratings['if']}]"

        clean_article = remove_doi_from_text(remove_naas_from_text(article))
        lines.append(f"{clean_article}{rating_str}")

        if link:
            lines.append(f"   CrossRef: {link}")

        lines.append("")

    return "\n".join(lines).strip()


def read_docx(file_content: bytes) -> str:
    """Read DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        paragraphs = [
            para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
        return ""


def main():
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>Publication Deduplicator</h1>
        <p>Deduplicate articles • Generate CrossRef links • Preserve ratings</p>
    </div>
    """, unsafe_allow_html=True)

    # Sidebar
    with st.sidebar:
        st.markdown("### Settings")

        similarity_threshold = st.slider(
            "Similarity Threshold",
            min_value=0.50,
            max_value=1.00,
            value=0.85,
            step=0.05,
            help="Higher = stricter matching"
        )

        trust_numbering = st.checkbox(
            "Trust article numbering",
            value=True,
            help="If checked, uses numbered format (1., 2., etc.) to identify articles"
        )

        show_crossref = st.checkbox(
            "Show CrossRef links",
            value=True,
            help="Display clickable DOI links"
        )

        auto_fetch_doi = st.checkbox(
            "Auto-fetch missing DOI from Crossref",
            value=True,
            help="Attempts to find DOI for articles without DOI (slower)"
        )

        st.markdown("---")

        st.markdown("### Features")
        st.markdown("""
        - **CrossRef Links**: Click to view article
        - **NAAS/IF Badges**: Color-coded ratings
        - **Format Detection**: Auto-detects numbered vs paragraph
        - **Verification Mode**: Compare detected vs expected count
        """)

    # Main content
    col1, col2 = st.columns([2, 1])

    with col1:
        st.markdown("### Input Articles")

        input_method = st.radio(
            "Input method:",
            ["Upload DOCX", "Paste Text"],
            horizontal=True,
            label_visibility="collapsed"
        )

        input_text = ""

        if input_method == "Upload DOCX":
            uploaded_file = st.file_uploader("Upload DOCX", type=['docx'])

            if uploaded_file:
                input_text = read_docx(uploaded_file.getvalue())
                if input_text:
                    st.success(f"Loaded: {uploaded_file.name}")
        else:
            input_text = st.text_area(
                "Paste articles",
                height=300,
                placeholder="""Paste your articles here...

For best results, use numbered format:
1. Author (Year). Title. Journal, Volume, Pages. DOI
2. Author (Year). Title. Journal, Volume, Pages. DOI
..."""
            )

    with col2:
        st.markdown("### Preview")
        if input_text:
            articles, format_type, info = parse_articles(
                input_text, trust_numbering)

            st.metric("Articles Detected", len(articles))
            st.caption(f"Format: **{format_type}**")

            if format_type == "numbered" and info.get('max_number'):
                expected = info['max_number']
                actual = len(articles)

                if expected != actual:
                    st.warning(
                        f"Expected {expected} (max #), found {actual}")
                    if info.get('missing_numbers'):
                        st.caption(
                            f"Missing: {info['missing_numbers'][:10]}...")
                else:
                    st.success(f"Count matches numbering")

            # Count articles with DOIs
            doi_count = sum(1 for a in articles if extract_doi(a))
            st.metric("With DOI (CrossRef)", doi_count)

            # Count ratings
            naas_count = sum(
                1 for a in articles if extract_rating(a).get('naas'))
            if naas_count:
                st.metric("With NAAS Rating", naas_count)
        else:
            st.info("Upload or paste articles")

    # Process button
    st.markdown("---")
    _, btn_col, _ = st.columns([1, 2, 1])
    with btn_col:
        process = st.button("Process Articles", type="primary",
                            use_container_width=True, disabled=not input_text)

    # Processing
    if process and input_text:
        with st.spinner("Analyzing..."):
            articles, format_type, info = parse_articles(
                input_text, trust_numbering)

            if not articles:
                st.error("No articles found. Check format.")
                return

            unique_articles, duplicate_groups = find_duplicates(
                articles, similarity_threshold)

            st.session_state['unique_articles'] = unique_articles
            st.session_state['all_articles'] = articles
            st.session_state['duplicate_groups'] = duplicate_groups
            st.session_state['parse_info'] = info
            st.session_state['processed'] = True

    # Results
    if st.session_state.get('processed', False):
        unique_articles = st.session_state['unique_articles']
        all_articles = st.session_state['all_articles']
        duplicate_groups = st.session_state['duplicate_groups']
        parse_info = st.session_state.get('parse_info', {})

        total = len(all_articles)
        unique = len(unique_articles)
        duplicates = total - unique
        reduction = (duplicates / total * 100) if total > 0 else 0

        # Metrics
        st.markdown(
            '<div class="section-header"><h2>Summary</h2></div>', unsafe_allow_html=True)

        cols = st.columns(4)
        with cols[0]:
            st.markdown(
                f'<div class="metric-card total"><p class="metric-value">{total}</p><p class="metric-label">Total</p></div>', unsafe_allow_html=True)
        with cols[1]:
            st.markdown(
                f'<div class="metric-card unique"><p class="metric-value">{unique}</p><p class="metric-label">Unique</p></div>', unsafe_allow_html=True)
        with cols[2]:
            st.markdown(
                f'<div class="metric-card duplicate"><p class="metric-value">{duplicates}</p><p class="metric-label">Duplicates</p></div>', unsafe_allow_html=True)
        with cols[3]:
            st.markdown(
                f'<div class="metric-card reduction"><p class="metric-value">{reduction:.1f}%</p><p class="metric-label">Reduction</p></div>', unsafe_allow_html=True)

        # Parsing verification
        if parse_info.get('format') == 'numbered' and parse_info.get('max_number'):
            expected = parse_info['max_number']
            if expected != total:
                st.markdown(f"""
                <div class="info-box warning">
                    <strong>Count Mismatch:</strong> Your document has articles numbered up to {expected}, 
                    but the parser found {total} articles. Missing numbers: {parse_info.get('missing_numbers', [])[:10]}
                </div>
                """, unsafe_allow_html=True)

        # Duplicates
        if duplicate_groups:
            st.markdown(
                '<div class="section-header"><h2>🔄 Duplicates</h2></div>', unsafe_allow_html=True)

            for idx, (first_idx, dup_indices) in enumerate(duplicate_groups.items(), 1):
                count = 1 + len(dup_indices)

                with st.expander(f"Group {idx} — {count} occurrences", expanded=idx <= 3):
                    st.markdown("**Kept:**")
                    html = format_article_html(
                        all_articles[first_idx],
                        show_crossref=show_crossref,
                        auto_fetch_doi=auto_fetch_doi
                    )

                    st.markdown(
                        f'<div class="article-card">{html[:600]}{"..." if len(all_articles[first_idx]) > 600 else ""}</div>', unsafe_allow_html=True)

                    st.markdown("**Removed:**")
                    for dup_idx in dup_indices:
                        sim = calculate_similarity(normalize_text(
                            all_articles[first_idx]), normalize_text(all_articles[dup_idx]))
                        st.markdown(
                            f'<div class="article-card duplicate"><strong>Position {dup_idx + 1}</strong> ({sim:.0%} similar)<br/>{all_articles[dup_idx][:200]}...</div>', unsafe_allow_html=True)
        else:
            st.markdown(
                '<div class="info-box success"><strong>No duplicates!</strong></div>', unsafe_allow_html=True)

        # Unique articles
        st.markdown(
            '<div class="section-header"><h2>Unique Articles</h2></div>', unsafe_allow_html=True)

        # Downloads
        dl_cols = st.columns(3)
        with dl_cols[0]:
            st.download_button("Download TXT", create_text_output(
                unique_articles), "unique_articles.txt", "text/plain", use_container_width=True)
        with dl_cols[1]:
            docx = create_docx(unique_articles)
            if docx:
                st.download_button("Download DOCX", docx, "unique_articles.docx",
                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        with dl_cols[2]:
            if st.button("Copy Mode", use_container_width=True):
                st.session_state['show_copy'] = True

        if st.session_state.get('show_copy'):
            st.text_area(
                "Copy:",
                create_text_output(unique_articles),
                height=300
            )

        # List articles
        with st.expander(f"All {unique} Unique Articles", expanded=False):
            for idx, article in enumerate(unique_articles, 1):
                html = format_article_html(
                    article,
                    show_crossref=show_crossref,
                    auto_fetch_doi=auto_fetch_doi
                )

                st.markdown(
                    f'<div class="article-card"><strong>{idx}.</strong> {html}</div>', unsafe_allow_html=True)

        # Frequency table
        st.markdown(
            '<div class="section-header"><h2>Frequency</h2></div>', unsafe_allow_html=True)

        import pandas as pd

        freq_data = []
        seen = set()

        for idx, article in enumerate(all_articles):
            h = generate_article_hash(article)
            if h in seen:
                continue
            seen.add(h)

            count = 1
            for first_idx, dups in duplicate_groups.items():
                if idx == first_idx:
                    count = 1 + len(dups)
                    break

            ratings = extract_rating(article)
            doi = extract_doi(article)

            freq_data.append({
                "Article": article[:70] + "...",
                "Count": count,
                "Rating": f"NAAS {ratings['naas']}" if 'naas' in ratings else (f"IF {ratings['if']}" if 'if' in ratings else ""),
                "DOI": "✅" if doi else "❌",
                "Status": "Unique" if count == 1 else f"Dup ({count}x)"
            })

        freq_data.sort(key=lambda x: x["Count"], reverse=True)
        df = pd.DataFrame(freq_data)

        filters = st.columns(3)
        with filters[0]:
            only_dups = st.checkbox("Only duplicates")
        with filters[1]:
            only_rated = st.checkbox("Only with ratings")
        with filters[2]:
            only_doi = st.checkbox("Only with DOI")

        if only_dups:
            df = df[df["Count"] > 1]
        if only_rated:
            df = df[df["Rating"] != ""]
        if only_doi:
            df = df[df["DOI"] == "✅"]

        st.dataframe(df, use_container_width=True, hide_index=True)

    # Sidebar footer
    st.sidebar.markdown("---")
    st.sidebar.markdown(
        """
        <small>
        <b>Publication Deduplicator v1.0</b><br>
        Developed by<br>
        <a href="https://scholar.google.com/citations?user=Es-kJk4AAAAJ&hl=en" target="_blank">
            Dr. Sandip Garai
        </a>
        &nbsp;·&nbsp;
        <a href="https://scholar.google.com/citations?user=0dQ7Sf8AAAAJ&hl=en&oi=ao" target="_blank">
            Dr. Kanaka K K
        </a><br>
        <a href="mailto:drgaraislab@gmail.com">Contact</a>
        </small>
        """,
        unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()

